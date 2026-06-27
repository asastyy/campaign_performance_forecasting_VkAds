from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_DIR = Path(__file__).resolve().parents[1]
OUTPUT_PATH = PROJECT_DIR / "docs" / "Сопроводительная записка НИР Сергеева АВ VK Ads.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_table_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_paragraph(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.10


def add_results_table(doc: Document) -> None:
    rows = [
        ["Модель", "Интерпретация", "Метрика, %"],
        ["Always zero", "Нулевой baseline: все доли равны 0", "445.86"],
        ["Month template shift 744", "Прогноз второго месяца по предыдущему месяцу без прямого replay", "11.31"],
        ["Month template + calibration", "Per-target log-bias calibration поверх frozen month-template", "11.09"],
        ["HGB temporal + embeddings", "Boosting baseline с user embedding features на expanding temporal split", "12.78"],
        ["Auction replay tie=0.5", "Replay открытой валидации по тем же часам history", "0.85"],
        ["Auction replay tie=0.483", "Validation-calibrated replay с подобранной tie-вероятностью", "0.84"],
    ]

    table = doc.add_table(rows=len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [Inches(1.85), Inches(3.95), Inches(0.95)]

    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx]
        keep_table_row_together(row)
        if row_idx == 0:
            set_repeat_table_header(row)
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.width = widths[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_idx == 0:
                set_cell_shading(cell, "F2F4F7")
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 2 else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(text)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            if row_idx == 0:
                run.bold = True


def build() -> None:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(
        "Сопроводительная записка к курсовой работе / НИР\n"
        "Декомпозиционное моделирование прогноза охвата рекламной кампании в аукционной системе VK"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle_run = subtitle.add_run(
        "Студентка: Сергеева Анастасия. Направление: машинное обучение и анализ данных. "
        "Предметная область: AdTech, прогнозирование reach/frequency, auction modeling."
    )
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor.from_string("555555")

    doc.add_heading("Актуальность", level=1)
    add_paragraph(
        doc,
        "Прогноз охвата рекламной кампании до запуска напрямую связан с медиапланированием: рекламодателю нужно заранее понимать, какая доля целевой аудитории увидит объявление хотя бы один, два или три раза при заданных CPM, площадках, периоде показа и составе аудитории. В аукционной рекламной системе эта оценка зависит не только от активности пользователя, но и от конкурентной ставки в конкретном показе, поэтому простая регрессия по агрегатам теряет часть предметной механики.",
    )
    add_paragraph(
        doc,
        "Задача VK хорошо соответствует современным AdTech-подходам: индустриальные forecasting-системы учитывают доступный инвентарь, таргетинг, длительность кампании, частотные ограничения и auction dynamics. В литературе близкая постановка рассматривается в работах о campaign performance forecasting и bid landscape modeling, где результат кампании собирается из последовательности аукционов.",
    )

    doc.add_heading("Обзор предметной области", level=1)
    add_paragraph(
        doc,
        "Предметная область естественно делится на три блока. Первый блок — supply, то есть оценка доступных opportunities: будет ли пользователь активен в нужные часы и на нужных площадках. Второй блок — auction: выиграет ли объявление с данным CPM относительно наблюдаемой максимальной ставки конкурентов. Третий блок — aggregation: перевод пользовательских показов и сессий в доли 1+ / 2+ / 3+.",
    )
    add_paragraph(
        doc,
        "В качестве основной исследовательской опоры использована статья Wang et al. (2024) “Know in AdVance”, где прогноз результата кампании формулируется через последовательность RTB-аукционов и совмещает auction-level и campaign-level представления. Дополнительно в обзоре рассмотрены AuctionNet как пример simulator/replay-среды, работы по censored bid landscape modeling и практические материалы VK ML / ads.",
    )

    doc.add_heading("Постановка задачи", level=1)
    add_paragraph(
        doc,
        "Входные данные: таблица пользователей `users.tsv`, история рекламных показов `history.tsv`, параметры кампаний `validate.tsv` и эталонные ответы `validate_answers.tsv`. Для каждой кампании известны CPM, временное окно, список площадок и список пользователей целевой аудитории.",
    )
    add_paragraph(
        doc,
        "Выход модели: три величины `at_least_one`, `at_least_two`, `at_least_three` — доли пользователей аудитории, которые увидят объявление хотя бы 1, 2 и 3 раза. Основная метрика — Smoothed Mean Log Accuracy Ratio с epsilon = 0.005; чем ближе значение к 0%, тем лучше.",
    )
    add_paragraph(
        doc,
        "Ограничение данных состоит в том, что в истории наблюдаются только уже выигранные показы и winning CPM, а полный лог проигранных аукционов недоступен. Поэтому auction-модуль реконструируется приближенно по observed winning price, а перенос в будущее делается через временной шаблон предыдущего месяца.",
    )

    doc.add_heading("Выбор метода", level=1)
    add_paragraph(
        doc,
        "После обсуждения с научным руководителем boosting выбран как сильный baseline на агрегированных признаках, а основная модель строится как интерпретируемая декомпозиция supply × auction × aggregation. AuctionNet оставлен в обзоре подходов как аргумент в пользу replay/simulation-валидации, но не как центральный метод реализации.",
    )
    add_paragraph(
        doc,
        "После экспертного комментария дизайн экспериментов скорректирован в сторону более промышленной постановки: supervised-baseline оценивается на expanding temporal split, калибровка выполняется отдельно для каждого из трех таргетов, а в feature engineering добавлены user embedding-like признаки. Embeddings рекламных item'ов не строятся, потому что в предоставленной истории нет идентификатора объявления (`ad_id`/`item_id`).",
    )
    add_number(doc, "Supply-модуль отбирает исторические opportunities по пользователю, часу и площадке.")
    add_number(doc, "Auction-модуль применяет правила задачи: если CPM кампании выше observed winning CPM, показ выигрывается; если равен — вероятность выигрыша 0.5.")
    add_number(doc, "Session-модуль ограничивает повторные показы: новая пользовательская сессия начинается при разрыве между показами не менее 6 часов.")
    add_number(doc, "Aggregation-модуль собирает вероятности по сессиям в пороговые метрики 1+ / 2+ / 3+ через Poisson-binomial динамическое программирование.")

    doc.add_heading("Исследование и воспроизводимость", level=1)
    add_paragraph(
        doc,
        "В ноутбуке реализованы и прогнаны несколько экспериментов: нулевой baseline, прогноз по шаблону предыдущего месяца (`source_shift = 744` часа), per-target calibration для month-template, replay открытой валидации по тем же часам, boosting baseline `HistGradientBoostingRegressor` на агрегированных признаках и multi-output MLP как прототип общей модели с тремя выходами. CatBoost/LightGBM можно использовать как замену baseline при наличии соответствующих библиотек; в текущей воспроизводимой версии выбран sklearn-бустинг без внешней установки.",
    )
    add_paragraph(
        doc,
        "Для embedding-like признаков пользователей строится компактное SVD-представление по историческим паттернам `publisher × hour-of-day × day-of-week`, после чего компоненты агрегируются по аудитории кампании. На expanding temporal split rolling per-target calibration улучшает month-template на будущих строках с 11.77% до 11.55%, а HGB с user embeddings дает 12.78%; multi-output MLP оказался слабее и не выбран финальным методом.",
    )
    add_paragraph(
        doc,
        "Код оформлен как небольшой проект: основной модуль `src/vk_ads_solution.py`, исполняемый ноутбук `VK_Ads_reach_forecasting_solution.ipynb`, runner `run_experiments.py` и сохраненные TSV-прогнозы в папке `outputs`. Путь к данным можно переопределить через переменную окружения `VK_ADS_DATA_DIR`.",
    )

    doc.add_heading("Метрики и результаты", level=1)
    add_paragraph(
        doc,
        "Основные результаты на предоставленном `validate.tsv` приведены ниже. Replay дает лучший результат, потому что открытые валидационные часы уже содержатся в `history.tsv`; последняя строка дополнительно калибрует вероятность равенства CPM по `validate_answers.tsv`. Per-target calibration немного улучшает month-template на открытой валидации, но для будущего скрытого теста корректнее ориентироваться на temporal split, rolling calibration, теоретическое правило tie=0.5 и развитие supply-модуля.",
    )
    add_results_table(doc)

    doc.add_heading("Итоги", level=1)
    add_paragraph(
        doc,
        "Получена рабочая модель, которая явно использует механику аукциона и сессий. На открытой валидации replay-вариант с теоретическим tie=0.5 достигает 0.85% по официальной метрике, а validation-calibrated вариант с tie=0.483 — 0.84%. Реалистичный перенос по предыдущему месяцу дает 11.31%, per-target calibration улучшает его до 11.09%, а rolling temporal calibration на будущих строках — до 11.55%. Это подтверждает, что декомпозиция не только интерпретируема, но и практически эффективна на данной постановке.",
    )
    add_paragraph(
        doc,
        "Ограничения работы: неполная наблюдаемость проигранных аукционов, зависимость от сезонности прошлого месяца и отсутствие полного production-лога с `ad_id`. Следующие шаги: обучить отдельный supply-модуль на временных и пользовательских признаках, заменить empirical replay на learned bid landscape / survival model, проверять устойчивость на rolling time-based holdout, а в production подтверждать эффект через A/B-тестирование.",
    )

    doc.add_heading("Состав подготовленного решения", level=1)
    add_bullet(doc, "Исполняемый ноутбук с экспериментами и сохраненными outputs.")
    add_bullet(doc, "Модуль с реализацией метрики, sessionization, replay-прогноза и feature engineering.")
    add_bullet(doc, "TSV-файлы с прогнозами для validate, включая финальный `validation_predictions_final.tsv`.")
    add_bullet(doc, "CSV с temporal-метриками экспертных корректировок: `expert_temporal_metrics.csv`.")
    add_bullet(doc, "Сопроводительная записка в формате `.docx`.")
    add_bullet(doc, "Презентация для защиты в формате `.pptx`.")

    doc.add_heading("Ключевые источники", level=1)
    add_paragraph(doc, "Wang et al. (2024). Know in AdVance: Linear-Complexity Forecasting of Ad Campaign Performance with Evolving User Interest. KDD 2024.")
    add_paragraph(doc, "Li et al. (2022). Arbitrary Distribution Modeling with Censorship in Real-Time Bidding Advertising. KDD 2022.")
    add_paragraph(doc, "Su et al. (2024). AuctionNet: A Novel Benchmark for Decision-Making in Large-Scale Games. NeurIPS 2024.")
    add_paragraph(doc, "VK ML / ads: материалы о рекламе, RTB и forecasting распределения ставок.")
    add_paragraph(doc, "Google Display & Video 360 Help: Create a plan for a campaign.")

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build()
