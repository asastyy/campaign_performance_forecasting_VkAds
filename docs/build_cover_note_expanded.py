from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUTPUT_PATH = (
    Path(__file__).resolve().parent
    / "Сопроводительная записка НИР Сергеева АВ VK Ads расширенная.docx"
)


def set_run_font(run, size: int = 12, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(30, 30, 30)


def add_paragraph(doc: Document, text: str = "", size: int = 12) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    set_run_font(run, size=size)


def add_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=16, bold=True)


def add_subheading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, size=13, bold=True)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    set_run_font(run, size=12)


def add_numbered(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    set_run_font(run, size=12)


def add_formula(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.name = "Cambria Math"
    run.font.size = Pt(11)
    run.font.bold = False
    run.font.color.rgb = RGBColor(20, 20, 20)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        element = tc_mar.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def write_cell(cell, text: str, *, bold: bool = False, size: int = 9, align_center: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    if align_center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.autofit = False

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Cm(widths_cm[idx])
        set_cell_shading(cell, "F2F4F7")
        write_cell(cell, header, bold=True, size=9, align_center=idx in (0, len(headers) - 1))

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].width = Cm(widths_cm[idx])
            write_cell(cells[idx], value, size=9, align_center=idx in (0, len(row) - 1))

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def add_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(18)
    run = paragraph.add_run(text)
    set_run_font(run, size=18, bold=True)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(12)
    styles["List Bullet"].font.name = "Arial"
    styles["List Bullet"].font.size = Pt(12)
    styles["List Number"].font.name = "Arial"
    styles["List Number"].font.size = Pt(12)

    add_title(
        doc,
        "Прогнозирование охвата и частоты рекламной кампании "
        "в аукционной системе VK Ads",
    )

    add_paragraph(
        doc,
        "Тематика программы: машинное обучение, анализ данных, рекламные технологии.",
    )
    add_paragraph(
        doc,
        "Направление: campaign performance forecasting и reach/frequency forecasting "
        "в аукционных рекламных системах.",
    )

    add_heading(doc, "Актуальность")
    add_paragraph(
        doc,
        "Задача прогнозирования охвата рекламной кампании важна для рекламных платформ "
        "и рекламодателей, поскольку до запуска кампании необходимо оценивать не только "
        "ожидаемый объём показов, но и распределение контактов внутри целевой аудитории.",
    )
    add_paragraph(
        doc,
        "Один и тот же суммарный объём показов может соответствовать разным сценариям: "
        "много пользователей увидели рекламу один раз или небольшая группа пользователей "
        "получила несколько повторных контактов. Поэтому метрики 1+, 2+ и 3+ являются "
        "практически значимыми для планирования охвата и частоты.",
    )
    add_paragraph(
        doc,
        "Задача усложняется тем, что показы формируются через рекламный аукцион. "
        "На результат влияют ставка CPM, время запуска кампании, площадки, активность "
        "пользователей, доступный рекламный инвентарь и правило пользовательских сессий.",
    )

    add_heading(doc, "Обзор предметной области")
    add_paragraph(
        doc,
        "Работа находится на пересечении нескольких направлений: прогнозирования "
        "эффективности рекламных кампаний, оценки охвата и частоты, а также моделирования "
        "рекламных аукционов.",
    )
    add_subheading(doc, "1. Campaign performance forecasting")
    add_paragraph(
        doc,
        "В современных работах по campaign performance forecasting, например AdVance, "
        "основное внимание уделяется прогнозированию агрегированных показателей кампании: "
        "стоимости, числа показов, кликов и конверсий. Такие модели используют богатые "
        "внутренние логи рекламных платформ: последовательности аукционов, клики, "
        "конверсии, candidate ads и поведенческие признаки пользователей.",
    )
    add_paragraph(
        doc,
        "Отличие данной работы состоит в другой целевой переменной: прогнозируются не "
        "агрегированные KPI кампании, а threshold reach - доли пользователей, получивших "
        "не менее 1, 2 и 3 показов.",
    )

    add_subheading(doc, "2. Reach/frequency forecasting и frequency capping")
    add_paragraph(
        doc,
        "В работах по reach measurement и frequency capping изучаются охват, частота "
        "контактов и ограничения на число показов одному пользователю. Это близко к "
        "рассматриваемой задаче по смыслу, однако чаще формулируется как задача "
        "измерения, оптимизации или ограничения частоты, а не как прогноз будущей "
        "кампании по ограниченным открытым аукционным логам.",
    )

    add_subheading(doc, "3. Auction simulation")
    add_paragraph(
        doc,
        "Работы по auction simulation, включая AuctionNet-подобные подходы, рассматривают "
        "моделирование рекламных аукционов и поведения bidding-стратегий. Они полезны "
        "как методологическая опора, но решают другую задачу: симуляцию или оптимизацию "
        "ставок, а не прогноз долей пользователей с 1+, 2+ и 3+ контактами.",
    )

    add_subheading(doc, "Итог обзора")
    add_paragraph(
        doc,
        "Близкие направления существуют, но точная публичная постановка с прогнозом "
        "P(N >= 1), P(N >= 2), P(N >= 3) при ограниченных VK-like auction logs встречается "
        "редко. Поэтому вклад работы состоит в адаптации идей прогнозирования рекламных "
        "кампаний к задаче threshold reach forecasting с жёстким контролем data leakage.",
    )
    add_paragraph(
        doc,
        "Для фиксации новизны работа сопоставлялась не только по названию темы, но и по "
        "типу данных, целевой переменной и воспроизводимости постановки.",
    )
    add_table(
        doc,
        ["Направление", "Что решает", "Связь с работой", "Отличие данной задачи"],
        [
            [
                "AdVance / campaign forecasting",
                "Прогноз KPI кампаний: показы, клики, конверсии, стоимость.",
                "Показывает актуальность прогнозирования рекламной эффективности.",
                "Цель здесь не KPI кампании, а доли пользователей с 1+, 2+ и 3+ контактами; в открытых данных нет кликов, креативов и полных auction logs.",
            ],
            [
                "Reach measurement и frequency capping",
                "Оценка охвата, частоты и ограничение повторных контактов.",
                "Близкая предметная логика: важно распределение контактов по пользователям.",
                "Чаще решается измерение или оптимизация частоты, а не прогноз будущей кампании по ограниченной истории показов.",
            ],
            [
                "Auction simulation / AuctionNet",
                "Симуляция аукциона, bidding-стратегий и рекламной среды.",
                "Даёт методологическую опору для учёта правил аукциона.",
                "Фокус на симуляции и ставках; итоговая цель не threshold reach по заданной аудитории.",
            ],
            [
                "Данная работа",
                "Прогноз at_least_one, at_least_two, at_least_three для будущей кампании.",
                "Объединяет историю показов, CPM-аукцион, сессии и агрегацию вероятностей.",
                "Постановка сделана для ограниченных VK-like логов и строгой temporal validation без data leakage.",
            ],
        ],
        [3.0, 4.1, 4.4, 5.2],
    )

    add_heading(doc, "Постановка задачи")
    add_paragraph(
        doc,
        "На вход подаются параметры будущей рекламной кампании: CPM, час начала и "
        "окончания, список площадок, размер аудитории и идентификаторы пользователей. "
        "Дополнительно доступны история показов и базовые пользовательские признаки.",
    )
    add_paragraph(
        doc,
        "На выходе необходимо предсказать три величины:",
    )
    add_bullet(doc, "at_least_one - доля пользователей, увидевших объявление хотя бы один раз;")
    add_bullet(doc, "at_least_two - доля пользователей, увидевших объявление хотя бы два раза;")
    add_bullet(doc, "at_least_three - доля пользователей, увидевших объявление хотя бы три раза.")
    add_paragraph(
        doc,
        "Формально для пользователя u из аудитории A рассматривается случайная величина "
        "N_u - число показов объявления этому пользователю. Тогда целевые значения "
        "задаются как средние вероятности по аудитории:",
    )
    add_formula(doc, "at_least_k = (1 / |A|) * sum_{u in A} P(N_u >= k),   k in {1, 2, 3}")

    add_heading(doc, "Метод")
    add_paragraph(
        doc,
        "В качестве основной модели выбран декомпозиционный replay-подход. Его идея "
        "состоит в том, чтобы не прогнозировать итоговый охват напрямую, а отдельно "
        "смоделировать исторический рекламный инвентарь, правила аукциона, пользовательские "
        "сессии и агрегацию вероятностей.",
    )
    add_subheading(doc, "Правило выигрыша аукциона")
    add_paragraph(
        doc,
        "Для каждого исторического события сравнивается CPM прогнозируемой кампании "
        "с winning CPM в истории. Вероятность выигрыша задаётся правилами задачи:",
    )
    add_formula(doc, "P(win) = 1,   если cpm_campaign > cpm_history")
    add_formula(doc, "P(win) = 0.5, если cpm_campaign = cpm_history")
    add_formula(doc, "P(win) = 0,   если cpm_campaign < cpm_history")

    add_subheading(doc, "Учет пользовательских сессий")
    add_paragraph(
        doc,
        "Система не показывает пользователю одно и то же объявление повторно внутри "
        "одной сессии. Новая сессия начинается после шести часов без показов. Поэтому "
        "показы сначала агрегируются на уровне пользовательских сессий, а затем для "
        "каждого пользователя вычисляется распределение числа выигранных сессий.",
    )
    add_formula(doc, "P_u(0) = product_{s in S_u} (1 - p_s)")
    add_formula(doc, "P_u(N >= 1) = 1 - P_u(0)")
    add_formula(doc, "P_u(N >= 2) = 1 - P_u(0) - P_u(1)")
    add_formula(doc, "P_u(N >= 3) = 1 - P_u(0) - P_u(1) - P_u(2)")
    add_paragraph(
        doc,
        "Значения P_u(1) и P_u(2) вычисляются через Poisson-binomial динамическую "
        "агрегацию по вероятностям пользовательских сессий.",
    )

    add_subheading(doc, "Исторические окна и blend")
    add_paragraph(
        doc,
        "Чтобы учесть разные временные масштабы, используются несколько past-only "
        "исторических компонент: monthly, daily и weekly. Итоговый прогноз получается "
        "геометрическим blend в лог-пространстве, что согласовано с логарифмической "
        "метрикой качества.",
    )
    add_formula(
        doc,
        "y_hat = exp(w_m log(y_m + eps) + w_d log(y_d + eps) + w_w log(y_w + eps)) - eps",
    )
    add_paragraph(
        doc,
        "Финальная конфигурация: daily_lags = 8, weekly_lags = 5, веса "
        "monthly/daily/weekly = 0.05 / 0.40 / 0.55, bias-калибровка отключена.",
    )

    add_heading(doc, "Защита от data leakage")
    add_paragraph(
        doc,
        "Для временных рекламных данных особенно важно не использовать информацию из "
        "будущего. В работе реализован строгий validation protocol:",
    )
    add_bullet(doc, "temporal split по времени запуска кампаний;")
    add_bullet(doc, "использование только past-only признаков и окон;")
    add_bullet(doc, "запрет на replay целевых часов validate;")
    add_bullet(doc, "отдельный locked final holdout;")
    add_bullet(doc, "фиксация прогноза и SHA-256 перед расчётом финальной метрики.")
    add_paragraph(
        doc,
        "Открытая validation-выборка частично пересекается с history по времени, поэтому "
        "наивное использование history до каждого hour_start дало бы более низкую метрику, "
        "но не соответствовало бы hidden-test постановке. В финальной оценке используется "
        "строгий test-like режим: история режется до первого часа validate.",
    )
    add_table(
        doc,
        ["Потенциальный риск", "Как закрыт в работе", "Зачем это важно"],
        [
            [
                "Использование будущих показов",
                "История ограничена cutoff до начала validate; признаки строятся только из прошлого.",
                "Иначе модель видит будущий рекламный инвентарь и получает завышенную метрику.",
            ],
            [
                "Подбор весов на финальном тесте",
                "Выбор конфигурации выполнен на development/calibration/pretest; final holdout закрыт до финального расчёта.",
                "Финальная оценка остаётся проверкой обобщения, а не результатом ручной подгонки.",
            ],
            [
                "Повторные показы внутри сессии",
                "История агрегируется в пользовательские сессии с правилом 6 часов.",
                "Без этого частота контактов переоценивается, особенно для активных пользователей.",
            ],
            [
                "Сравнение с моделями из статей",
                "Проводится концептуально по постановке, данным и целям, а не как прямой численный benchmark.",
                "В статьях используются другие закрытые данные и другие целевые переменные.",
            ],
        ],
        [4.1, 6.5, 5.6],
    )

    add_heading(doc, "Базовые модели и эксперименты")
    add_paragraph(
        doc,
        "Эксперименты были выстроены как последовательное усиление baseline: от простого "
        "исторического шаблона к декомпозиции по временным окнам, затем к калибровке и "
        "сравнению с ML-моделями. Такой дизайн позволил не только получить лучший скор, "
        "но и объяснить, почему финальная модель является более устойчивой.",
    )
    add_paragraph(doc, "Ключевые проверенные варианты:")
    add_bullet(doc, "monthly replay baseline - прогноз по похожему окну прошлого месяца;")
    add_bullet(doc, "daily replay - несколько последних дневных окон для учёта коротких изменений спроса;")
    add_bullet(doc, "weekly replay - недельные лаги для сезонности по дням недели и часам;")
    add_bullet(doc, "geometric blend - объединение monthly/daily/weekly компонент в лог-пространстве;")
    add_bullet(doc, "target-wise diagnostics - отдельный подбор компонент для 1+, 2+ и 3+ как проверка чувствительности таргетов;")
    add_bullet(doc, "градиентный бустинг на агрегированных признаках пользователей, кампании и истории;")
    add_bullet(doc, "бустинг с replay/decomposition features как сильный ML-baseline;")
    add_bullet(doc, "нейросетевая multi-task постановка с общим представлением и отдельными головами для трёх таргетов.")
    add_subheading(doc, "Эксперименты с replay-декомпозицией")
    add_table(
        doc,
        ["Этап", "Проверка", "Метрика отбора", "Final holdout", "Вывод"],
        [
            [
                "1",
                "Monthly replay baseline.",
                "11.77% temporal; 11.55% с rolling calibration",
                "11.70%",
                "Честный baseline, но плохо ловит краткосрочные изменения инвентаря.",
            ],
            [
                "2",
                "Daily replay: несколько последних дневных окон.",
                "9.53% purged OOF",
                "не финальная модель",
                "Улучшает baseline за счёт актуального спроса и свежей активности.",
            ],
            [
                "3",
                "Weekly replay: похожие дни недели и часы.",
                "9.87% purged OOF",
                "не финальная модель",
                "Добавляет недельную сезонность, но отдельно слабее blend.",
            ],
            [
                "4",
                "Geometric blend monthly/daily/weekly.",
                "9.46% purged OOF; 8.03% calibration; 8.17% pretest",
                "9.54%",
                "Финальная модель: лучшая устойчивость на закрытом temporal holdout.",
            ],
            [
                "5",
                "Target-wise blend и past-median calibration.",
                "9.43% purged OOF; target log-error diagnostics",
                "не финальная модель",
                "Полезно как диагностика, но финально выбран более простой устойчивый scalar blend.",
            ],
        ],
        [1.0, 5.1, 3.7, 2.3, 5.0],
    )
    doc.add_page_break()
    add_subheading(doc, "ML-baselines и проверка усложнения модели")
    add_table(
        doc,
        ["Этап", "Проверка", "Метрика отбора", "Final holdout", "Вывод"],
        [
            [
                "6",
                "Gradient boosting на user/history/campaign aggregates.",
                "12.78% temporal для HGB; 16.33-27.67% в строгих вариантах",
                "30.89%",
                "Без replay-логики табличная модель плохо переносится во времени.",
            ],
            [
                "7",
                "Boosting + replay/decomposition features.",
                "6.61% calibration; 6.92% pretest",
                "10.94%",
                "Сильный pretest, но хуже final holdout: признак переобучения на режим валидации.",
            ],
            [
                "8",
                "Multi-task MLP с общим представлением и головами для 1+, 2+ и 3+.",
                "38.12% temporal",
                "не финальная модель",
                "Близко к современным статьям концептуально, но в текущих ограниченных логах нестабильно.",
            ],
        ],
        [1.0, 5.1, 3.7, 2.3, 5.0],
    )
    add_paragraph(
        doc,
        "Главный вывод из экспериментов: простая ML-модель на агрегатах не заменяет "
        "механистическую декомпозицию. Для этой постановки критично явно учитывать "
        "правило выигрыша CPM-аукциона, сессионность и temporal split. Поэтому финальная "
        "модель выбрана не только по минимальной промежуточной метрике, но и по устойчивости "
        "между calibration, pretest и locked final holdout.",
    )

    doc.add_page_break()
    add_heading(doc, "Метрика качества")
    add_paragraph(
        doc,
        "Использовалась официальная метрика задачи - Smoothed Mean Log Accuracy Ratio. "
        "Она оценивает относительное расхождение между предсказанными и фактическими "
        "долями пользователей по трём таргетам.",
    )
    add_formula(
        doc,
        "Score = 100% * ( exp( (1 / (3n)) * sum_{i=1..n} sum_{j=1..3} "
        "|log((Pred_{ij} + eps) / (Actual_{ij} + eps))| ) - 1 )",
    )
    add_paragraph(
        doc,
        "Здесь eps = 0.005. Сглаживание нужно, чтобы метрика была устойчивой при "
        "малых фактических долях. Чем ниже значение метрики, тем лучше качество прогноза.",
    )

    add_heading(doc, "Результаты")
    add_paragraph(
        doc,
        "Итоговое сравнение проводилось на locked final holdout. Метрика минимизируется: "
        "чем ниже значение, тем точнее прогноз долей 1+, 2+ и 3+.",
    )
    add_table(
        doc,
        ["Модель", "Что использует", "Final metric", "Интерпретация"],
        [
            [
                "Monthly baseline",
                "Исторический шаблон прошлого месяца.",
                "11.70%",
                "Честная нижняя планка качества для сезонного replay-подхода.",
            ],
            [
                "Boosting + replay features",
                "190 агрегированных и replay-признаков; обучение на pretest.",
                "10.94%",
                "Сильный ML-baseline, но хуже переносится на финальный будущий период.",
            ],
            [
                "Static/history boosting",
                "64 агрегированных признака без полноценной replay-декомпозиции.",
                "30.89%",
                "Показывает, что одних табличных агрегатов недостаточно.",
            ],
            [
                "Финальная декомпозиционная модель",
                "CPM-аукцион, сессии, daily/weekly/monthly replay и геометрический blend.",
                "9.54%",
                "Лучший и наиболее устойчивый результат; улучшение monthly baseline примерно на 18%.",
            ],
        ],
        [4.2, 6.0, 2.4, 4.2],
    )
    add_paragraph(
        doc,
        "Итоговая модель улучшила monthly baseline примерно на 18% относительно и "
        "примерно на 13% относительно лучшего boosting baseline. Важный практический вывод: "
        "в этой задаче качество даёт не усложнение модели само по себе, а правильная "
        "декомпозиция механизма показа рекламы.",
    )

    add_heading(doc, "Почему не использовалась предобученная модель")
    add_paragraph(
        doc,
        "Использование внешней предобученной модели в данной постановке не является "
        "методологически обоснованным. В датасете отсутствуют тексты объявлений, "
        "изображения креативов, клики, конверсии, идентификаторы объявлений и полные "
        "auction logs. Поэтому невозможно корректно применить универсальную pretrained "
        "модель или воспроизвести модели из статей, обученные на закрытых логах рекламных "
        "платформ.",
    )
    add_paragraph(
        doc,
        "Вместо этого выбран интерпретируемый подход, основанный на механизме формирования "
        "показов: истории рекламного инвентаря, правилах CPM-аукциона и пользовательских "
        "сессиях.",
    )

    add_heading(doc, "Выводы")
    add_paragraph(
        doc,
        "В работе построена воспроизводимая модель прогнозирования охвата и частоты "
        "рекламной кампании в аукционной системе. Модель явно учитывает механизм "
        "формирования показов и защищена от утечек из будущего.",
    )
    add_paragraph(
        doc,
        "Академический вклад работы состоит в адаптации идей campaign performance "
        "forecasting к задаче threshold reach forecasting при ограниченных аукционных "
        "логах. Практический результат - улучшение качества прогноза относительно "
        "исторического baseline и ML-baseline на строгой temporal validation.",
    )

    add_heading(doc, "Ограничения и перспективы")
    add_paragraph(
        doc,
        "Основные ограничения связаны с неполнотой данных: отсутствуют проигравшие ставки, "
        "клики, конверсии, признаки креативов и идентификаторы объявлений. Поэтому прямое "
        "численное сравнение с современными статьями, такими как AdVance, невозможно; "
        "сравнение проводится концептуально по постановке задачи, данным и подходам.",
    )
    add_paragraph(doc, "Дальнейшее развитие работы может включать:")
    add_numbered(doc, "добавление признаков объявлений, креативов, кликов и конверсий при наличии данных;")
    add_numbered(doc, "проверку multi-task нейросетевой модели с общим представлением и отдельными головами для 1+, 2+ и 3+;")
    add_numbered(doc, "отдельную калибровку вероятностей для каждого таргета;")
    add_numbered(doc, "сравнение с моделями на полных auction logs при доступе к закрытым логам рекламной платформы.")

    add_heading(doc, "Использованные источники")
    add_bullet(doc, "Wang et al. Know in AdVance: Linear-Complexity Forecasting of Ad Campaign Performance with Evolving User Interest, 2024. https://arxiv.org/abs/2405.10681")
    add_bullet(doc, "Gao, Qiao. Reach Measurement, Optimization and Frequency Capping in Online Advertising, 2025. https://arxiv.org/abs/2501.04882")
    add_bullet(doc, "AuctionNet: benchmarking and simulation approaches for advertising auctions, 2024. https://arxiv.org/abs/2412.10798")
    add_bullet(doc, "Arrate et al. Large-scale analysis of user exposure to online advertising in Facebook, 2018. https://arxiv.org/abs/1811.10921")

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build()
